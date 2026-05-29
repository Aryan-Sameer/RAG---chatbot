"""
ingest.py - Multi-format folder ingestion (PDF, DOCX, XLSX, MD, TXT)
"""
import os
import json
from datetime import datetime, timezone
import pandas as pd
from docx import Document
from pypdf import PdfReader
from langchain_core.documents import Document as LCDocument
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_experimental.text_splitter import SemanticChunker
from db_connection import get_chroma_client

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".md", ".txt"}
COLLECTION_NAME = "documents"
MANIFEST_FILE = ".ingest_manifest.json"
EMBEDDING_MODEL_NAME = "nomic-ai/nomic-embed-text-v1.5"


def _build_embedding_model():
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_NAME,
        model_kwargs={"device": "cpu", "trust_remote_code": True},
    )

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif ext == '.docx':
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif ext in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path)
            text = df.to_string()
        elif ext in ['.md', '.txt']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        return text
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return None

def _list_supported_files(folder_path):
    files = []
    for filename in sorted(os.listdir(folder_path)):
        file_path = os.path.join(folder_path, filename)
        if not os.path.isfile(file_path):
            continue
        ext = os.path.splitext(filename)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue
        files.append(file_path)
    return files


def _file_signature(file_path):
    stat = os.stat(file_path)
    return {
        "size": stat.st_size,
        "mtime_ns": stat.st_mtime_ns
    }


def _manifest_path(db_path):
    return os.path.join(db_path, MANIFEST_FILE)


def _build_current_manifest(files):
    return {
        os.path.basename(path): _file_signature(path)
        for path in files
    }


def _read_previous_manifest(db_path):
    manifest_path = _manifest_path(db_path)
    if not os.path.exists(manifest_path):
        return {}
    try:
        with open(manifest_path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        return payload.get("files", {})
    except Exception:
        return {}


def _write_manifest(db_path, files_manifest):
    os.makedirs(db_path, exist_ok=True)
    payload = {
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "files": files_manifest,
    }
    with open(_manifest_path(db_path), "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2)


def save_uploaded_files(folder_path, uploads):
    """Save uploaded files to the knowledge-base folder.

    uploads: iterable of (filename, file_bytes)
    Returns list of saved file metadata dicts.
    """
    os.makedirs(folder_path, exist_ok=True)
    saved = []
    errors = []

    for filename, content in uploads:
        safe_name = os.path.basename(filename)
        ext = os.path.splitext(safe_name)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            errors.append(f"{safe_name}: unsupported type (allowed: {', '.join(sorted(SUPPORTED_EXTENSIONS))})")
            continue
        dest = os.path.join(folder_path, safe_name)
        with open(dest, "wb") as f:
            f.write(content)
        stat = os.stat(dest)
        saved.append({
            "name": safe_name,
            "size_bytes": stat.st_size,
            "modified_at": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
            "extension": ext,
        })

    return saved, errors


def list_knowledge_base_files(folder_path="./database/context"):
    if not os.path.exists(folder_path):
        return []
    records = []
    for file_path in _list_supported_files(folder_path):
        stat = os.stat(file_path)
        records.append({
            "name": os.path.basename(file_path),
            "size_bytes": stat.st_size,
            "modified_at": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
            "extension": os.path.splitext(file_path)[1].lower(),
        })
    return records


def run_folder_ingestion(folder_path="./database/context", db_path="./database/chroma_db", force_rebuild=False):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
        print(f"Created {folder_path} folder. Add your files there!")
        return {
            "success": True,
            "reingested": False,
            "changed": False,
            "message": "Knowledge-base folder was missing and is now created.",
            "files_processed": 0,
            "chunks_ingested": 0,
        }

    files = _list_supported_files(folder_path)
    current_manifest = _build_current_manifest(files)
    previous_manifest = _read_previous_manifest(db_path)
    changed = current_manifest != previous_manifest

    if not force_rebuild and not changed:
        print("✓ No file changes detected. Skipping ingestion.")
        return {
            "success": True,
            "reingested": False,
            "changed": False,
            "message": "No changes in knowledge-base files. Ingestion skipped.",
            "files_processed": 0,
            "chunks_ingested": 0,
        }

    embeddings = _build_embedding_model()
    client = get_chroma_client(db_path)

    if force_rebuild or changed:
        try:
            client.delete_collection(COLLECTION_NAME)
            print(f"✓ Deleted existing collection '{COLLECTION_NAME}' for rebuild.")
        except Exception:
            # Collection might not exist yet, which is fine.
            pass

    vectorstore = Chroma(
        collection_name=COLLECTION_NAME,
        client=client,
        embedding_function=embeddings,
    )
    splitter = SemanticChunker(
        embeddings,
        breakpoint_threshold_type="percentile"
    )

    files_processed = 0
    chunks_ingested = 0

    for file_path in files:
        filename = os.path.basename(file_path)
        print(f"→ Processing {filename}...")
        raw_text = extract_text(file_path)
        if raw_text:
            chunks = splitter.split_text(raw_text)
            if not chunks:
                continue
            ids = [f"{filename}_{i}" for i in range(len(chunks))]
            metadatas = [{"source": filename, "chunk_index": i} for i in range(len(chunks))]
            documents = [
                LCDocument(page_content=chunk, metadata=metadata)
                for chunk, metadata in zip(chunks, metadatas)
            ]
            vectorstore.add_documents(documents=documents, ids=ids)
            print(f"✓ Ingested {filename}")
            files_processed += 1
            chunks_ingested += len(chunks)

    _write_manifest(db_path, current_manifest)
    return {
        "success": True,
        "reingested": True,
        "changed": changed,
        "message": f"Ingestion complete. Processed {files_processed} file(s), {chunks_ingested} chunk(s).",
        "files_processed": files_processed,
        "chunks_ingested": chunks_ingested,
    }

if __name__ == "__main__":
    result = run_folder_ingestion()
    print(result)
